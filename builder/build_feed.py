#!/usr/bin/env python3
"""Article Quotes feed builder.

Reads the Cowork article archive (one PDF per article), extracts the
title / author / source / URL / summary / high-impact points / notable
quotes, and writes:

  1. A static feed under docs/feed/ (served by GitHub Pages) that the
     Android app downloads:
        manifest.json          shard list with hashes (app syncs only changed shards)
        quotes/<YYYY-MM>.json  quote pool sharded by archive month
        articles/<id>.json     full summary for one article
        index.json             lightweight article index for browsing/search
  2. One Word document per article in the Google Drive sync folder
     (Drive shows them as documents that open in Google Docs on the phone).

Incremental: a cache keyed on filename+mtime means a daily run only parses
new or changed PDFs.  Safe to re-run at any time.
"""
from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import os
import re
import sys
import time
from pathlib import Path

import fitz  # PyMuPDF

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
CONFIG_PATH = HERE / "config.json"

# ----------------------------------------------------------------------------
# Config
# ----------------------------------------------------------------------------

DEFAULT_CONFIG = {
    "archive_dir": r"C:\Claude\Claude Cowork\Projects\Article Archive - Education & Leadership",
    "drive_dir": r"C:\Users\losey\My Drive (timothyjamesmills@gmail.com)\Claude\Article Summaries",
    "feed_dir": str(ROOT / "docs" / "feed"),
    "cache_dir": str(HERE / ".cache"),
    # Categories (filename prefix before " - (") that feed the phone app.
    "feed_categories": [
        "Leadership",
        "Family",
        "Classical Education",
        "Education - Practice & Culture",
        "Education - Learning & Curriculum",
        "Education - Ed-Tech",
    ],
    # Categories written to Google Drive (all of them by default).
    "drive_categories": "all",
    "min_quote_chars": 40,
    "max_quote_chars": 600,
}


def load_config() -> dict:
    cfg = dict(DEFAULT_CONFIG)
    if CONFIG_PATH.exists():
        cfg.update(json.loads(CONFIG_PATH.read_text(encoding="utf-8")))
    return cfg


def long_path(p: str) -> Path:
    """Windows: allow paths longer than 260 chars (archive filenames are long)."""
    if os.name == "nt" and not p.startswith("\\\\?\\"):
        p = "\\\\?\\" + os.path.abspath(p)
    return Path(p)


# ----------------------------------------------------------------------------
# Parsing
# ----------------------------------------------------------------------------

FILENAME_RE = re.compile(r"^(?P<cat>.+?) - \((?P<auth>[^)]*)\) (?P<title>.+?) \((?P<m>\d{1,2})\.(?P<d>\d{1,2})\.(?P<y>\d{2})\)\.pdf$")
FILENAME_ISO_RE = re.compile(r"^(?P<cat>.+?) - \((?P<auth>[^)]*)\) (?P<title>.+?) \((?P<y>\d{4})\.(?P<m>\d{1,2})\.(?P<d>\d{1,2})\)\.pdf$")

KV_RE = re.compile(r"^(Title|Author|Source|Category|URL|Url|Date|Email date|Email Date|Archived|Published)\s*:\s*(.*)$", re.I)
SECTION_RE = re.compile(
    r"^(?P<name>summary|key takeaways|high[- ]impact points|key points|notable quotes|key quotes|selected quotes|quotes|full text|full article)\s*:?\s*$",
    re.I,
)
QUOTE_START_RE = re.compile(r"^(?:\[\d+\]\s*|\d+[.)]\s*|•\s*|-\s*)?[“\"‘']")
MARKER_ONLY_RE = re.compile(r"^(?:\[\d+\]|\d+[.)]|•|-)$")
OPEN_Q = "“\"‘"
CLOSE_Q = "”\"’"


def canonical_section(name: str) -> str:
    n = name.lower()
    if n.startswith("summary"):
        return "summary"
    if "quote" in n:
        return "quotes"
    if "full" in n:
        return "fulltext"
    return "bullets"


def clean_ws(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def strip_quote_marks(q: str) -> str:
    q = clean_ws(q)
    q = re.sub(r"^(?:\[\d+\]\s*|\d+[.)]\s*|•\s*|-\s*)", "", q)
    while q and q[0] in OPEN_Q:
        q = q[1:]
    while q and q[-1] in CLOSE_Q:
        q = q[:-1]
    return q.strip()


def parse_pdf(path: Path, cfg: dict) -> dict | None:
    m = FILENAME_RE.match(path.name) or FILENAME_ISO_RE.match(path.name)
    if not m:
        return None
    gd = m.groupdict()
    year = int(gd["y"]) if len(gd["y"]) == 4 else 2000 + int(gd["y"])
    try:
        fdate = dt.date(year, int(gd["m"]), int(gd["d"]))
    except ValueError:
        fdate = dt.date.fromtimestamp(path.stat().st_mtime)

    try:
        with open(path, "rb") as fh:
            data = fh.read()
        doc = fitz.open(stream=data, filetype="pdf")
        text = "".join(page.get_text() for page in doc)
        doc.close()
    except Exception as exc:  # corrupt PDF
        print(f"  ! cannot read {path.name}: {exc}", file=sys.stderr)
        return None

    lines = [ln.strip() for ln in text.splitlines()]
    fields: dict[str, str] = {}
    title_lines: list[str] = []
    sections: dict[str, list[str]] = {"summary": [], "bullets": [], "quotes": []}
    current: str | None = None  # None = header phase

    for ln in lines:
        if not ln:
            continue
        sec = SECTION_RE.match(ln)
        if sec:
            current = canonical_section(sec.group("name"))
            if current == "fulltext":
                break
            continue
        if current is None:
            kv = KV_RE.match(ln)
            if kv:
                key = kv.group(1).lower().replace(" ", "_")
                fields[key] = kv.group(2).strip()
            elif not fields.get("title"):
                # Free lines before/among the header are the title (may wrap).
                if len(title_lines) < 3 and len(ln) < 200:
                    title_lines.append(ln)
            continue
        if current == "fulltext":
            break
        sections[current].append(ln)

    title = fields.get("title") or clean_ws(" ".join(title_lines)) or gd["title"]
    title = clean_ws(title)
    author = clean_ws(fields.get("author", "")) or gd["auth"]
    author = author.strip("() ")
    if author.count("(") != author.count(")"):
        author = re.sub(r"\s*\([^)]*$", "", author).strip()
    if not author:
        author = gd["auth"]
    source = clean_ws(fields.get("source", ""))
    url = fields.get("url", "").strip()
    if url and not re.match(r"^https?://", url):
        url = ""
    date_display = fields.get("date") or fields.get("email_date") or ""
    # Date embedded at the end of the Source line ("..., May 16, 2026")
    if not date_display:
        dm = re.search(r"((?:January|February|March|April|May|June|July|August|September|October|November|December) \d{1,2}, \d{4})", source)
        if dm:
            date_display = dm.group(1)
    if not date_display:
        date_display = fdate.strftime("%B %d, %Y").replace(" 0", " ")

    summary = clean_ws(" ".join(sections["summary"]))

    # Bullets: a "•" may sit alone on a line, followed by wrapped text.
    bullets: list[str] = []
    buf: list[str] = []
    for ln in sections["bullets"]:
        starts = ln.startswith(("•", "- ", "– ")) or re.match(r"^\d+[.)]\s", ln) is not None
        if starts:
            if buf:
                bullets.append(clean_ws(" ".join(buf)))
            buf = [re.sub(r"^(•|- |– |\d+[.)]\s)\s*", "", ln)]
        else:
            buf.append(ln)
    if buf:
        bullets.append(clean_ws(" ".join(buf)))
    bullets = [b for b in bullets if len(b) > 15]

    # Quotes: start at an opening quote mark; end when a line ends with a closing mark.
    quotes: list[str] = []
    qbuf: list[str] = []
    for ln in sections["quotes"]:
        marker_only = MARKER_ONLY_RE.match(ln) is not None
        closed = bool(qbuf) and qbuf[-1].rstrip()[-1:] in CLOSE_Q
        if marker_only or (QUOTE_START_RE.match(ln) and (not qbuf or closed)):
            if qbuf:
                quotes.append(strip_quote_marks(" ".join(qbuf)))
            qbuf = [] if marker_only else [ln]
        else:
            qbuf.append(ln)
    if qbuf:
        quotes.append(strip_quote_marks(" ".join(qbuf)))
    lo, hi = cfg["min_quote_chars"], cfg["max_quote_chars"]
    seen = set()
    good_quotes = []
    for q in quotes:
        if lo <= len(q) <= hi and q.lower() not in seen:
            seen.add(q.lower())
            good_quotes.append(q)

    category = gd["cat"].strip()
    art_id = hashlib.sha1(path.name.encode("utf-8")).hexdigest()[:12]
    return {
        "id": art_id,
        "file": path.name,
        "category": category,
        "title": title,
        "author": author,
        "source": source,
        "date": fdate.isoformat(),
        "date_display": date_display,
        "url": url,
        "summary": summary,
        "bullets": bullets,
        "quotes": good_quotes,
    }


# ----------------------------------------------------------------------------
# Cache
# ----------------------------------------------------------------------------

def cache_key(path: Path) -> str:
    st = path.stat()
    return hashlib.sha1(f"{path.name}|{st.st_size}|{int(st.st_mtime)}".encode("utf-8")).hexdigest()


def load_cache(cache_dir: Path) -> dict:
    f = cache_dir / "parsed.json"
    if f.exists():
        try:
            return json.loads(f.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            pass
    return {}


def save_cache(cache_dir: Path, cache: dict) -> None:
    cache_dir.mkdir(parents=True, exist_ok=True)
    tmp = cache_dir / "parsed.json.tmp"
    tmp.write_text(json.dumps(cache, ensure_ascii=False), encoding="utf-8")
    tmp.replace(cache_dir / "parsed.json")


# ----------------------------------------------------------------------------
# Feed writing
# ----------------------------------------------------------------------------

def write_json(path: Path, data, compact=True) -> str:
    path.parent.mkdir(parents=True, exist_ok=True)
    txt = json.dumps(data, ensure_ascii=False, separators=(",", ":") if compact else (",", ": "), indent=None if compact else 2)
    old = path.read_text(encoding="utf-8") if path.exists() else None
    if old != txt:
        path.write_text(txt, encoding="utf-8")
    return hashlib.sha1(txt.encode("utf-8")).hexdigest()[:16]


def build_feed(articles: list[dict], feed_dir: Path, cfg: dict) -> dict:
    feed_cats = set(cfg["feed_categories"])
    feed_articles = [a for a in articles if a["category"] in feed_cats and a["quotes"]]
    feed_articles.sort(key=lambda a: (a["date"], a["title"]), reverse=True)

    # Per-article detail files
    art_dir = feed_dir / "articles"
    art_dir.mkdir(parents=True, exist_ok=True)
    keep = set()
    for a in feed_articles:
        keep.add(f"{a['id']}.json")
        write_json(art_dir / f"{a['id']}.json", {
            "id": a["id"], "category": a["category"], "title": a["title"], "author": a["author"],
            "source": a["source"], "date": a["date"], "dateDisplay": a["date_display"], "url": a["url"],
            "summary": a["summary"], "points": a["bullets"], "quotes": a["quotes"],
        })
    for stale in art_dir.glob("*.json"):
        if stale.name not in keep:
            stale.unlink()

    # Quote shards by archive month
    shards: dict[str, list] = {}
    for a in feed_articles:
        ym = a["date"][:7]
        shards.setdefault(ym, []).append({
            "id": a["id"], "c": a["category"], "t": a["title"], "a": a["author"], "d": a["date"],
            "q": a["quotes"],
        })
    q_dir = feed_dir / "quotes"
    q_dir.mkdir(parents=True, exist_ok=True)
    shard_list = []
    keep = set()
    for ym in sorted(shards):
        name = f"{ym}.json"
        keep.add(name)
        sha = write_json(q_dir / name, shards[ym])
        shard_list.append({"name": ym, "path": f"quotes/{name}", "sha": sha,
                           "articles": len(shards[ym]), "quotes": sum(len(x["q"]) for x in shards[ym])})
    for stale in q_dir.glob("*.json"):
        if stale.name not in keep:
            stale.unlink()

    # Index
    index = [{"id": a["id"], "c": a["category"], "t": a["title"], "a": a["author"], "d": a["date"], "n": len(a["quotes"])}
             for a in feed_articles]
    write_json(feed_dir / "index.json", index)

    cats = {}
    for a in feed_articles:
        cats[a["category"]] = cats.get(a["category"], 0) + 1
    manifest = {
        "version": 1,
        "generated": dt.datetime.now(dt.timezone.utc).isoformat(timespec="seconds"),
        "articleCount": len(feed_articles),
        "quoteCount": sum(len(a["quotes"]) for a in feed_articles),
        "categories": cats,
        "shards": shard_list,
    }
    write_json(feed_dir / "manifest.json", manifest, compact=False)
    return manifest


# ----------------------------------------------------------------------------
# Google Drive documents
# ----------------------------------------------------------------------------

def write_drive_docs(articles: list[dict], drive_dir: Path, cfg: dict) -> int:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("python-docx not installed; skipping Drive documents", file=sys.stderr)
        return 0
    if not Path(str(drive_dir).replace("\\\\?\\", "")).parent.exists():
        print(f"Drive folder not found ({drive_dir.parent}); skipping Drive documents", file=sys.stderr)
        return 0
    cats = cfg["drive_categories"]
    written = 0
    for a in articles:
        if cats != "all" and a["category"] not in cats:
            continue
        out_dir = drive_dir / a["category"]
        out_dir.mkdir(parents=True, exist_ok=True)
        out = out_dir / (Path(a["file"]).stem + ".docx")
        if out.exists():
            continue
        d = Document()
        style = d.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        d.add_heading(a["title"], level=1)
        meta = d.add_paragraph()
        meta.add_run("Author: ").bold = True
        meta.add_run(a["author"] + "\n")
        meta.add_run("Source: ").bold = True
        meta.add_run((a["source"] or "") + (f", {a['date_display']}" if a["date_display"] and a["date_display"] not in (a["source"] or "") else "") + "\n")
        meta.add_run("Category: ").bold = True
        meta.add_run(a["category"])
        if a["url"]:
            p = d.add_paragraph()
            p.add_run("URL: ").bold = True
            p.add_run(a["url"])
        d.add_heading("Summary", level=2)
        d.add_paragraph(a["summary"] or "(no summary extracted)")
        if a["bullets"]:
            d.add_heading("High-Impact Points", level=2)
            for b in a["bullets"]:
                d.add_paragraph(b, style="List Bullet")
        if a["quotes"]:
            d.add_heading("Notable Quotes", level=2)
            for i, q in enumerate(a["quotes"], 1):
                p = d.add_paragraph(style="List Number")
                r = p.add_run(f"“{q}”")
                r.italic = True
        tmp = out.with_suffix(".docx.tmp")
        d.save(tmp)
        tmp.replace(out)
        written += 1
    return written


# ----------------------------------------------------------------------------
# Main
# ----------------------------------------------------------------------------

def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--no-drive", action="store_true", help="skip writing Word documents to Google Drive")
    ap.add_argument("--reparse", action="store_true", help="ignore the parse cache")
    ap.add_argument("--limit", type=int, default=0, help="parse at most N PDFs (testing)")
    args = ap.parse_args()

    cfg = load_config()
    archive = long_path(cfg["archive_dir"])
    feed_dir = Path(cfg["feed_dir"])
    cache_dir = Path(cfg["cache_dir"])
    drive_dir = long_path(cfg["drive_dir"])
    if not archive.exists():
        print(f"Archive folder not found: {archive}", file=sys.stderr)
        return 2

    t0 = time.time()
    cache = {} if args.reparse else load_cache(cache_dir)
    pdfs = sorted(archive / n for n in os.listdir(archive)
                  if n.lower().endswith(".pdf") and not n.startswith(("_", ".", "~")))
    if args.limit:
        pdfs = pdfs[: args.limit]
    articles: list[dict] = []
    parsed_new = 0
    skipped = 0
    new_cache: dict = {}
    for p in pdfs:
        k = cache_key(p)
        if k in cache:
            new_cache[k] = cache[k]
        else:
            a = parse_pdf(p, cfg)
            parsed_new += 1
            if a is None:
                skipped += 1
                continue
            new_cache[k] = a
        articles.append(new_cache[k])
    save_cache(cache_dir, new_cache)

    manifest = build_feed(articles, feed_dir, cfg)
    drive_written = 0 if args.no_drive else write_drive_docs(articles, drive_dir, cfg)

    no_quotes = sum(1 for a in articles if not a["quotes"])
    no_summary = sum(1 for a in articles if not a["summary"])
    print(f"PDFs: {len(pdfs)}  parsed now: {parsed_new}  unparseable names: {skipped}")
    print(f"Articles without quotes: {no_quotes}   without summary: {no_summary}")
    print(f"Feed: {manifest['articleCount']} articles, {manifest['quoteCount']} quotes, {len(manifest['shards'])} shards")
    print("Categories in feed: " + ", ".join(f"{k} {v}" for k, v in sorted(manifest["categories"].items())))
    print(f"Drive documents written: {drive_written}")
    print(f"Done in {time.time() - t0:.1f}s")
    return 0


if __name__ == "__main__":
    sys.exit(main())
